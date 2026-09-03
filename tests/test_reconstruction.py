#!/usr/bin/env python3
"""Regression and causal-chain tests for the v0.3.0 pipeline."""

from __future__ import annotations

import copy
import json
import os
import shutil
import subprocess
import sys
import uuid
import hashlib
from pathlib import Path
from zipfile import ZipFile

import docx
import pytest
from PIL import Image, ImageDraw

REPO_ROOT = Path(__file__).resolve().parent.parent
SCRIPTS = REPO_ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

import build_docx
import check_consistency
from canonical_schema import docx_image_asset_ids, generation_policy, load_canonical_schema, versioned_asset
from manifest_schema import ManifestValidationError, asset_path_map, sha256_file, validate_manifest_file
from reconstruct_assets import MAX_INPUT_PIXELS, reconstruct_all_assets
from render_docx import audit_rendered_pages, find_soffice, render_docx
from run_pipeline import run_pipeline


class RecordingBackend:
    name = "test_fixture_generation_provider"
    mode = "image_to_image"
    model = "fixture-complete-artwork-model"

    def __init__(self, snapshots: dict[str, Path]):
        self.snapshots = snapshots
        self.requests: list[dict] = []

    def execute(self, request: dict, output_path: Path) -> dict:
        self.requests.append(copy.deepcopy(request))
        shutil.copyfile(self.snapshots[request["stage_id"]], output_path)
        return {"provider_request_id": f"test-{request['stage_id']}", "test_only": True}


def artwork(path: Path, accent: tuple[int, int, int] = (232, 168, 62), mode: str = "RGB", size=(320, 240)) -> Path:
    image = Image.new("RGB", size, (18, 38, 68))
    draw = ImageDraw.Draw(image)
    draw.rectangle((20, 20, size[0] - 20, size[1] - 20), outline=(126, 220, 178), width=6)
    draw.ellipse((size[0] // 3, size[1] // 4, 2 * size[0] // 3, 3 * size[1] // 4), fill=accent)
    draw.polygon([(0, size[1]), (size[0] // 2, size[1] // 2), (size[0], size[1])], fill=(35, 105, 116))
    if mode == "RGBA":
        image = image.convert("RGBA")
    elif mode == "L":
        image = image.convert("L")
    image.save(path)
    return path


def analysis(path: Path, *, max_iterations: int = 3) -> Path:
    path.write_text(json.dumps({
        "subject": "漂浮城市中的公共温室",
        "composition": "横向三角构图，主体位于中心偏上",
        "perspective": "低机位广角透视",
        "environment": "云海与模块化城市平台",
        "foreground": "近景植物叶片形成遮挡",
        "middle_ground": "温室与人物活动区域",
        "background_visual": "云层、远景平台与天空",
        "palette": "深蓝、青绿与暖金色",
        "lighting": "右上方暖色主光和冷色环境光",
        "atmosphere": "清晨、开放、具有空气透视",
        "visual_style": "完整的未来生态概念艺术",
        "material": "玻璃、金属、植物和云雾",
        "detail": "结构清晰，材质可辨，留有最终精修空间",
        "theme": "未来生态共同体",
        "max_iterations": max_iterations,
        "convergence_threshold": 0.98,
    }, ensure_ascii=False), encoding="utf-8")
    return path


def make_plan(root: Path, final: Path, count: int) -> dict[str, Path]:
    plan: dict[str, Path] = {}
    colors = [(175, 42, 55), (80, 60, 185), (35, 155, 80), (215, 115, 35), (80, 190, 205)]
    for version in range(1, count + 1):
        snapshot = root / f"fixture-v{version}.png"
        if version == count:
            shutil.copyfile(final, snapshot)
        else:
            artwork(snapshot, accent=colors[version - 1])
            with Image.open(snapshot) as image:
                altered = Image.new("RGB", image.size, colors[version - 1])
                altered.save(snapshot)
        plan[f"generation_v{version}"] = snapshot
    return plan


def execute_pipeline(root: Path, count: int = 2, mode: str = "RGB") -> tuple[RecordingBackend, Path]:
    root.mkdir(parents=True, exist_ok=True)
    final = artwork(root / "输入作品.png", mode=mode)
    backend = RecordingBackend(make_plan(root, final, count))
    manifest = run_pipeline(str(final), str(root / "输出"), "浮城温室", "创意测试赛", str(analysis(root / "分析.json", max_iterations=4)), backend=backend, max_iterations=4)
    return backend, manifest


@pytest.fixture
def workdir():
    root = REPO_ROOT / ".test-work" / str(uuid.uuid4())
    root.mkdir(parents=True)
    try:
        yield root
    finally:
        shutil.rmtree(root, ignore_errors=True)


@pytest.fixture(scope="module")
def pipeline_output() -> tuple[RecordingBackend, Path]:
    root = REPO_ROOT / ".test-work" / f"pipeline-{uuid.uuid4()}"
    root.mkdir(parents=True)
    result = execute_pipeline(root, 2)
    yield result
    shutil.rmtree(root, ignore_errors=True)


def test_v1_is_real_execution_and_complete(pipeline_output) -> None:
    backend, path = pipeline_output
    manifest = validate_manifest_file(path)
    record = manifest.generation_records[0]
    assert record.status == "success"
    assert record.output == versioned_asset("generation", 1)[1]
    assert record.backend == backend.name != "opencv_filter"
    assert record.complete_artwork is True
    with Image.open(asset_path_map(manifest, path)[record.stage_id]) as image:
        assert image.width >= 64 and image.height >= 64


def test_prompt_v1_precedes_and_enters_request(pipeline_output) -> None:
    backend, path = pipeline_output
    manifest = validate_manifest_file(path)
    record = manifest.generation_records[0]
    request = json.loads(asset_path_map(manifest, path)[record.request_asset_id].read_text(encoding="utf-8"))
    assert request["status"] == "requested"
    assert request["prompt"] == record.prompt == backend.requests[0]["prompt"]
    assert request["requested_at"] <= record.generated_at


def test_difference_analysis_reads_actual_v1_and_final(pipeline_output) -> None:
    _, path = pipeline_output
    manifest = validate_manifest_file(path)
    record = manifest.generation_records[0]
    difference = json.loads(asset_path_map(manifest, path)[record.difference_analysis_asset_id].read_text(encoding="utf-8"))
    assert difference["input_asset_ids"] == ["generation_v1", "final_artwork"]
    assert difference["metrics"]["overall_convergence"] < 0.98
    assert any(char.isdigit() for char in difference["composition"][0])


def test_prompt_v2_has_traceable_evolution(pipeline_output) -> None:
    _, path = pipeline_output
    manifest = validate_manifest_file(path)
    first, second = manifest.generation_records
    evolution = manifest.prompt_record[1].evolution
    assert second.prompt != first.prompt
    assert evolution.keep and evolution.modify and evolution.reason
    assert first.stage_id in second.input_assets
    assert manifest.prompt_record[1].source_difference_asset_id == first.difference_analysis_asset_id
    assert manifest.prompt_record[1].source_adjustment_reason_asset_id == first.adjustment_reason_asset_id
    adjustment = json.loads(asset_path_map(manifest, path)[first.adjustment_reason_asset_id].read_text(encoding="utf-8"))
    assert adjustment["source_difference_asset_id"] == first.difference_analysis_asset_id


def test_v2_is_a_second_real_execution(pipeline_output) -> None:
    backend, path = pipeline_output
    manifest = validate_manifest_file(path)
    assert len(backend.requests) == len(manifest.generation_records) == 2
    assert backend.requests[1]["stage_id"] == "generation_v2"
    assert manifest.generation_records[1].status == "success"
    paths = asset_path_map(manifest, path)
    with Image.open(paths["final_artwork"]) as final_image, Image.open(paths["generation_v1"]) as v1_image, Image.open(paths["generation_v2"]) as v2_image:
        assert v1_image.size == v2_image.size == final_image.size
    assert all(record.complete_artwork for record in manifest.generation_records)


@pytest.mark.parametrize("count", [1, 2, 3])
def test_dynamic_version_counts_and_stage_graph(workdir: Path, count: int) -> None:
    _, path = execute_pipeline(workdir / f"case-{count}", count)
    manifest = validate_manifest_file(path)
    assert len(manifest.generation_records) == count
    assert len([stage for stage in manifest.stage_graph if stage.kind == "generation"]) == count
    assert len(manifest.stage_graph) == 2 * count + 2
    assert manifest.stage_graph[-1].inputs[0].asset_id == f"generation_v{count}"


def test_all_outputs_share_execution_record_source(pipeline_output) -> None:
    _, path = pipeline_output
    manifest = validate_manifest_file(path)
    paths = asset_path_map(manifest, path)
    doc = docx.Document(paths["statement_docx"])
    doc_text = "\n".join([*(p.text for p in doc.paragraphs), *(cell.text for table in doc.tables for row in table.rows for cell in row.cells)])
    for record in manifest.generation_records:
        assert json.loads(paths[record.record_asset_id].read_text(encoding="utf-8")) == record.model_dump(mode="json")
        assert record.prompt in paths["prompt_record"].read_text(encoding="utf-8")
        assert record.backend in paths["parameter_record"].read_text(encoding="utf-8")
        assert record.record_asset_id in json.dumps([stage.model_dump(mode="json") for stage in manifest.stage_graph], ensure_ascii=False)
        assert record.model in doc_text


def test_docx_media_matches_dynamic_schema_and_records(pipeline_output) -> None:
    _, path = pipeline_output
    manifest = validate_manifest_file(path)
    paths = asset_path_map(manifest, path)
    expected_ids = ["reconstructed_sketch", "reconstructed_lineart", "reconstructed_color_block", *[record.stage_id for record in manifest.generation_records], "final_artwork"]
    document = docx.Document(paths["statement_docx"])
    actual_ids = [shape._inline.docPr.get("descr") for shape in document.inline_shapes]
    assert actual_ids == expected_ids
    actual_hashes = []
    for shape in document.inline_shapes:
        rel_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        actual_hashes.append(hashlib.sha256(document.part.related_parts[rel_id].blob).hexdigest())
    assert actual_hashes == [sha256_file(paths[asset_id]) for asset_id in expected_ids]
    assert set(docx_image_asset_ids(generation_versions=2)) == {"final_artwork", "reconstructed_sketch", "reconstructed_lineart", "reconstructed_color_block", "generation_v1", "generation_v2"}


def test_no_generation_filters_remain() -> None:
    source = (SCRIPTS / "reconstruct_assets.py").read_text(encoding="utf-8")
    assert "visual_study" not in source
    assert "addWeighted" not in source
    assert "generate_visual" not in source


def test_backend_unavailable_is_explicit(workdir: Path) -> None:
    final = artwork(workdir / "final.png")
    env = os.environ.copy()
    for key in list(env):
        if key.startswith("AIGC_IMAGE_GENERATION_"):
            env.pop(key)
    process = subprocess.run([sys.executable, str(SCRIPTS / "run_pipeline.py"), "--input", str(final), "--output-dir", str(workdir / "out"), "--title", "测试", "--competition", "测试赛", "--analysis-json", str(analysis(workdir / "analysis.json"))], capture_output=True, text=True, encoding="utf-8", errors="replace", env=env)
    assert process.returncode == 3
    assert "generation_backend_unavailable" in process.stderr
    assert not list((workdir / "out").glob("*_generation_v*.png"))


@pytest.mark.parametrize("mutation", ["empty", "missing_field", "empty_graph", "duplicate_stage", "same_prompt", "bad_backend", "extra_field", "missing_image", "zero_image", "corrupt_image", "bad_hash", "absolute_path", "record_drift", "request_drift", "difference_drift"])
def test_invalid_manifests_fail_nonzero(pipeline_output, workdir: Path, mutation: str) -> None:
    _, source = pipeline_output
    root = workdir / mutation
    shutil.copytree(source.parent, root)
    path = root / source.name
    data = json.loads(path.read_text(encoding="utf-8"))
    paths = {item["id"]: root / item["path"] for item in data["assets"]}
    if mutation == "empty":
        path.write_bytes(b"")
    elif mutation == "missing_field":
        data.pop("generation_records")
    elif mutation == "empty_graph":
        data["stage_graph"] = []
    elif mutation == "duplicate_stage":
        data["stage_graph"].append(copy.deepcopy(data["stage_graph"][0]))
    elif mutation == "same_prompt":
        data["generation_records"][1]["prompt"] = data["generation_records"][0]["prompt"]
        data["prompt_record"][1]["prompt"] = data["generation_records"][0]["prompt"]
    elif mutation == "bad_backend":
        data["generation_records"][0]["backend"] = "opencv_filter"
        data["parameter_record"][0]["backend"] = "opencv_filter"
    elif mutation == "extra_field":
        data["assets"][0]["unexpected_field"] = "x"
    elif mutation in {"missing_image", "zero_image", "corrupt_image", "bad_hash", "absolute_path"}:
        asset = next(item for item in data["assets"] if item["id"] == "generation_v1")
        image_path = paths["generation_v1"]
        if mutation == "missing_image": image_path.unlink()
        elif mutation == "zero_image": image_path.write_bytes(b"")
        elif mutation == "corrupt_image":
            image_path.write_bytes(b"broken")
            asset.update(size_bytes=6, sha256=sha256_file(image_path))
        elif mutation == "bad_hash": asset["sha256"] = "0" * 64
        elif mutation == "absolute_path": asset["path"] = str(image_path.resolve())
    elif mutation == "record_drift":
        record_path = paths["generation_record_v1"]
        record = json.loads(record_path.read_text(encoding="utf-8")); record["model"] = "tampered"; record_path.write_text(json.dumps(record), encoding="utf-8")
        asset = next(item for item in data["assets"] if item["id"] == "generation_record_v1"); asset.update(size_bytes=record_path.stat().st_size, sha256=sha256_file(record_path))
    elif mutation == "request_drift":
        request_path = paths["generation_request_v1"]
        request = json.loads(request_path.read_text(encoding="utf-8")); request["prompt"] = "tampered"; request_path.write_text(json.dumps(request), encoding="utf-8")
        asset = next(item for item in data["assets"] if item["id"] == "generation_request_v1"); asset.update(size_bytes=request_path.stat().st_size, sha256=sha256_file(request_path))
    elif mutation == "difference_drift":
        diff_path = paths["difference_analysis_v1"]
        diff = json.loads(diff_path.read_text(encoding="utf-8")); diff["input_asset_ids"] = ["generation_v1"]; diff_path.write_text(json.dumps(diff), encoding="utf-8")
        asset = next(item for item in data["assets"] if item["id"] == "difference_analysis_v1"); asset.update(size_bytes=diff_path.stat().st_size, sha256=sha256_file(diff_path))
    if mutation != "empty": path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
    process = subprocess.run([sys.executable, str(SCRIPTS / "validate_manifest.py"), str(path)], capture_output=True, text=True, encoding="utf-8", errors="replace")
    assert process.returncode != 0


@pytest.mark.parametrize("mode", ["RGB", "RGBA", "L"])
def test_unicode_path_and_image_modes(workdir: Path, mode: str) -> None:
    source = artwork(workdir / f"中文 路径 {mode}.png", mode=mode)
    result = reconstruct_all_assets(str(source), str(workdir / "输出 资产"))
    assert set(result) == {"reconstructed_sketch", "reconstructed_lineart", "reconstructed_color_block"}
    for value in result.values():
        with Image.open(value) as image: image.verify()


def test_oversized_input_fails(workdir: Path) -> None:
    source = workdir / "huge.png"
    Image.new("1", (8000, 7000)).save(source)
    assert 8000 * 7000 > MAX_INPUT_PIXELS
    with pytest.raises(ValueError, match="exceeds"):
        reconstruct_all_assets(str(source), str(workdir / "out"))


def test_real_temporary_install_includes_runtime(workdir: Path) -> None:
    shell = shutil.which("pwsh") or shutil.which("powershell")
    if os.name == "nt":
        assert shell
        command = [shell, "-NoProfile", "-ExecutionPolicy", "Bypass", "-File", str(SCRIPTS / "install.ps1"), "-Platform", "codex", "-DestinationRoot", str(workdir), "-SkipFontInstall"]
    else:
        command = ["bash", str(SCRIPTS / "install.sh"), "codex", "--destination-root", str(workdir), "--skip-font-install"]
    process = subprocess.run(command, capture_output=True, text=True, encoding="utf-8", errors="replace")
    assert process.returncode == 0, process.stderr
    installed = workdir / "aigc-competition-statement"
    for relative in ("scripts/image_generation_backend.py", "scripts/analyze_generation_difference.py", "schema/canonical-assets.yaml", "requirements.txt"):
        assert (installed / relative).is_file()


def test_build_docx_rejects_missing_pre_generation_image(pipeline_output, workdir: Path) -> None:
    _, source = pipeline_output
    root = workdir / "missing-docx-image"
    shutil.copytree(source.parent, root)
    path = root / source.name
    manifest = validate_manifest_file(path)
    paths = asset_path_map(manifest, path)
    paths["reconstructed_sketch"].unlink()
    with pytest.raises(ManifestValidationError, match="does not exist"):
        build_docx.build_docx_from_manifest(str(path), str(paths["statement_docx"]))


def test_cli_files_have_git_executable_mode() -> None:
    git = shutil.which("git")
    assert git
    cli_files = {
        "scripts/analyze_generation_difference.py", "scripts/build_docx.py", "scripts/check_consistency.py",
        "scripts/image_generation_backend.py", "scripts/install.sh", "scripts/reconstruct_assets.py",
        "scripts/render_docx.py", "scripts/run_pipeline.py", "scripts/scan_placeholders.py",
        "scripts/validate_manifest.py", "tests/fixture_generation_backend.py", "tests/run_e2e_verification.py",
    }
    process = subprocess.run([git, "ls-files", "-s", *sorted(cli_files)], cwd=REPO_ROOT, capture_output=True, text=True)
    assert process.returncode == 0, process.stderr
    modes = {line.split()[3]: line.split()[0] for line in process.stdout.splitlines()}
    assert modes == {path: "100755" for path in cli_files}


def test_schema_policy_and_repository_consistency() -> None:
    assert load_canonical_schema()["manifest_version"] == "0.3.1"
    assert generation_policy()["absolute_max_versions"] > generation_policy()["minimum_versions"]
    assert check_consistency.check_version_sync() == []
    assert check_consistency.check_canonical_schema() == []
    assert check_consistency.check_no_utf8_bom() == []
    assert check_consistency.check_font_bundle() == []


def test_libreoffice_renders_every_page(pipeline_output, workdir: Path) -> None:
    if find_soffice() is None:
        if os.environ.get("REQUIRE_LIBREOFFICE") == "1": pytest.fail("LibreOffice is required but soffice was not found")
        pytest.skip("LibreOffice not installed on this host")
    _, path = pipeline_output
    manifest = validate_manifest_file(path)
    pages = render_docx(asset_path_map(manifest, path)["statement_docx"], workdir / "render")
    audit_rendered_pages(pages)
    assert pages
