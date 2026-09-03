#!/usr/bin/env python3
"""Build and structurally verify an A4 competition statement DOCX."""

from __future__ import annotations

import argparse
import hashlib
import json
import sys
from pathlib import Path

import docx
from PIL import Image
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt

from canonical_schema import docx_image_asset_ids
from manifest_schema import asset_path_map, validate_manifest_file

FONT_NAME = "Noto Sans SC"
TABLE_WIDTH_DXA = 9024

for stream in (sys.stdout, sys.stderr):
    if stream and hasattr(stream, "reconfigure"):
        stream.reconfigure(encoding="utf-8", errors="replace")


def _set_run_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), FONT_NAME)


def _keep(paragraph, *, next_paragraph: bool = False, lines: bool = True) -> None:
    properties = paragraph._p.get_or_add_pPr()
    if next_paragraph:
        properties.append(OxmlElement("w:keepNext"))
    if lines:
        properties.append(OxmlElement("w:keepLines"))


def _add_heading(document, text: str, level: int) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    _keep(paragraph, next_paragraph=True)


def _add_body(document, text: str, *, label: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.25
    if label:
        _set_run_font(paragraph.add_run(label), bold=True)
    _set_run_font(paragraph.add_run(text))


def _set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA} DXA: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "0")
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        for cell, value in zip(row.cells, widths):
            cell.width = Inches(value / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), str(value))


def _fill_table(table, rows: list[list[str]], widths: list[int]) -> None:
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    _set_run_font(run, size=8.5)
    _set_table_geometry(table, widths)
    for cell in table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            _set_run_font(run, size=9, bold=True)


def _add_image_with_caption(document, image_path: Path, caption: str, asset_id: str) -> None:
    with Image.open(image_path) as image:
        width_px, height_px = image.size
    max_width, max_height = 5.75, 6.65
    ratio = min(max_width / width_px, max_height / height_px)
    width, height = width_px * ratio, height_px * ratio
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(6)
    image_paragraph.paragraph_format.space_after = Pt(2)
    _keep(image_paragraph, next_paragraph=True)
    shape = image_paragraph.add_run().add_picture(str(image_path), width=Inches(width), height=Inches(height))
    shape._inline.docPr.set("descr", asset_id)
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(8)
    _keep(caption_paragraph)
    _set_run_font(caption_paragraph.add_run(caption), size=9)


def _configure_document(document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    for level, size in ((1, 15), (2, 13), (3, 11)):
        style = styles[f"Heading {level}"]
        style.font.name = FONT_NAME
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        style.paragraph_format.space_after = Pt(6 if level == 1 else 4)


def _verify_docx(docx_path: Path, image_assets: list[tuple[str, Path]], version_count: int) -> None:
    result = docx.Document(docx_path)
    if len(result.inline_shapes) != len(image_assets):
        raise AssertionError(f"DOCX image count mismatch: expected {len(image_assets)}, got {len(result.inline_shapes)}")
    embedded = []
    for shape in result.inline_shapes:
        asset_id = shape._inline.docPr.get("descr")
        relationship_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        blob = result.part.related_parts[relationship_id].blob
        embedded.append((asset_id, hashlib.sha256(blob).hexdigest()))
    expected = [(asset_id, hashlib.sha256(path.read_bytes()).hexdigest()) for asset_id, path in image_assets]
    if embedded != expected:
        raise AssertionError("DOCX drawings do not correspond one-for-one with canonical image assets")
    text = "\n".join(paragraph.text for paragraph in result.paragraphs)
    required_text = ["实际工具", "版权状态", "完整作品连续版本", *[f"Prompt V{number}" for number in range(1, version_count + 1)], *[f"Generation V{number}" for number in range(1, version_count + 1)]]
    missing = [token for token in required_text if token not in text and not any(token in cell.text for table in result.tables for row in table.rows for cell in row.cells)]
    if missing:
        raise AssertionError(f"DOCX did not render required records: {missing}")
    forbidden = ["{作品名称}", "{赛事名称}", "待补齐", "待插入", "TODO", "TBD", "PLACEHOLDER"]
    leaked = [token for token in forbidden if token in text]
    if leaked:
        raise AssertionError(f"DOCX contains placeholders: {leaked}")


def build_docx_from_manifest(manifest_path: str, output_docx_path: str) -> str:
    manifest = validate_manifest_file(manifest_path, allow_missing_asset_ids={"statement_docx"})
    paths = asset_path_map(manifest, manifest_path)
    output = Path(output_docx_path).resolve()
    if output != paths["statement_docx"]:
        raise ValueError(f"Output path must match statement_docx manifest asset: {paths['statement_docx']}")
    image_ids = docx_image_asset_ids(generation_versions=len(manifest.generation_records))
    missing = [asset_id for asset_id in image_ids if not paths[asset_id].is_file()]
    if missing:
        raise FileNotFoundError(f"Canonical images missing: {missing}")

    document = docx.Document()
    _configure_document(document)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)
    _keep(title, next_paragraph=True)
    _set_run_font(title.add_run(f"【{manifest.artwork.title}】AIGC 创作过程说明书"), size=18, bold=True)

    _add_heading(document, "一、作品基本信息", 1)
    for label, value in (
        ("作品名称：", manifest.artwork.title), ("参赛赛事：", manifest.artwork.competition),
        ("作品类型：", manifest.artwork.type), ("创作主题：", manifest.artwork.theme),
        ("技术路径：", manifest.artwork.pipeline),
    ):
        _add_body(document, value, label=label)

    _add_heading(document, "二、创作构思与立意", 1)
    _add_body(document, manifest.creative_rationale.background, label="创作背景：")
    _add_body(document, manifest.creative_rationale.visual_concept, label="视觉构思：")
    _add_body(document, manifest.creative_rationale.ai_collaboration, label="工具协同：")

    _add_heading(document, "三、Dynamic Stage Graph", 1)
    for index, stage in enumerate(manifest.stage_graph, start=1):
        _add_heading(document, f"3.{index} {stage.title}", 2)
        _add_body(document, stage.purpose, label="目的：")
        if stage.source_record_asset_id:
            _add_body(document, stage.source_record_asset_id, label="事实来源：")

    _add_heading(document, "四、前期视觉设计 / 输入素材", 1)
    caption_by_id = {
        output_item.asset_id: output_item.label
        for stage in manifest.stage_graph
        for output_item in stage.outputs
    }
    caption_by_id["final_artwork"] = "最终作品（用户提供文件）"
    preliminary_ids = [asset_id for asset_id in image_ids if asset_id.startswith("reconstructed_")]
    figure_number = 1
    for asset_id in preliminary_ids:
        evidence = next(asset.evidence_level for asset in manifest.assets if asset.id == asset_id)
        _add_image_with_caption(document, paths[asset_id], f"图 {figure_number} {caption_by_id.get(asset_id, asset_id)} {evidence}（非 Generation Version）", asset_id)
        figure_number += 1

    _add_heading(document, "五、AIGC 完整作品连续版本", 1)
    _add_body(document, "Generation V1/V2/V3…均为同一幅作品的完整画面快照，不是人物、背景或局部资产。")
    for index, (record, prompt_item, parameter_item) in enumerate(zip(manifest.generation_records, manifest.prompt_record, manifest.parameter_record), start=1):
        _add_heading(document, f"5.{index} Generation V{index}", 2)
        _add_body(document, "、".join(record.input_assets), label="输入素材：")
        _add_body(document, record.prompt, label=f"Prompt V{index}：")
        if prompt_item.source_difference_asset_id:
            _add_body(document, f"{prompt_item.source_difference_asset_id} + {prompt_item.source_adjustment_reason_asset_id}", label="Prompt Evolution 来源：")
        _add_body(document, record.backend, label="实际工具：")
        _add_body(document, record.model, label="实际模型：")
        parameter_table = document.add_table(rows=1, cols=2)
        parameter_table.rows[0].cells[0].text = "实际参数"
        parameter_table.rows[0].cells[1].text = "值"
        _fill_table(parameter_table, [[key, str(value)] for key, value in parameter_item.parameters.items()], [1872, 7152])
        _add_body(document, "；".join(prompt_item.evolution.keep) or "首轮建立方向", label="KEEP：")
        _add_body(document, "；".join(prompt_item.evolution.modify) or "首轮建立方向", label="MODIFY：")
        _add_body(document, "；".join(prompt_item.evolution.add) or "无", label="ADD：")
        _add_body(document, "；".join(prompt_item.evolution.reduce) or "无", label="REDUCE：")
        difference = json.loads(paths[record.difference_analysis_asset_id].read_text(encoding="utf-8"))
        adjustment = json.loads(paths[record.adjustment_reason_asset_id].read_text(encoding="utf-8"))
        _add_body(document, "；".join(difference["priority_adjustments"]), label=f"V{index} 实际问题：")
        _add_body(document, "；".join(item["adjustment"] for item in adjustment["items"]), label="修改原因与动作：")
        evidence = next(asset.evidence_level for asset in manifest.assets if asset.id == record.stage_id)
        _add_image_with_caption(document, paths[record.stage_id], f"图 {figure_number} Generation V{index}：同一作品完整版本 {evidence}", record.stage_id)
        figure_number += 1

    _add_heading(document, "六、Final Artwork", 1)
    _add_body(document, "Final 与最后一个 Generation Version 保持主体、构图与风格继承关系，承担最后精修或真实后期处理。")
    final_evidence = next(asset.evidence_level for asset in manifest.assets if asset.id == "final_artwork")
    _add_image_with_caption(document, paths["final_artwork"], f"图 {figure_number} Final Artwork {final_evidence}", "final_artwork")

    _add_heading(document, "七、版权、原创性与原始工具状态", 1)
    for label, claim in (
        ("版权状态：", manifest.provenance.copyright), ("原创性状态：", manifest.provenance.originality),
        ("原始工具状态：", manifest.provenance.original_tool),
    ):
        source = f"；确认来源：{claim.confirmation_source}" if claim.confirmation_source else ""
        _add_body(document, f"{claim.value} {claim.evidence_level}{source}", label=label)
    _add_body(document, manifest.disclaimer)

    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    if not output.is_file() or output.stat().st_size == 0:
        raise AssertionError(f"DOCX was not created: {output}")
    expected_order = [*preliminary_ids, *[record.stage_id for record in manifest.generation_records], "final_artwork"]
    _verify_docx(output, [(asset_id, paths[asset_id]) for asset_id in expected_order], len(manifest.generation_records))
    return str(output)


def main() -> int:
    parser = argparse.ArgumentParser(description="Build AIGC competition statement DOCX")
    parser.add_argument("--manifest", "-m", required=True)
    parser.add_argument("--output", "-o", required=True)
    args = parser.parse_args()
    output = build_docx_from_manifest(args.manifest, args.output)
    print(f"DOCX generated and verified: {output} ({Path(output).stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
