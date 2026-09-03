#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
scan_placeholders.py - Scans rendered artifacts for leaked placeholders.

Scans rendered markdown, json, and docx files to ensure zero dangling placeholders
such as "{作品名称}", "待补齐", "待确认", "TODO", etc. remain in the final deliverables.
"""

import os
import sys
if sys.stdout:
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except Exception:
        pass
import argparse
from pathlib import Path
import docx

FORBIDDEN_PATTERNS = [
    "{作品名称}",
    "{赛事名称}",
    "{提示词内容}",
    "{值}",
    "待补齐",
    "待确认",
    "待插入",
    "TODO",
    "TBD",
    "PLACEHOLDER"
]


def scan_text(text: str, source_name: str) -> list:
    """Returns list of found forbidden patterns in text."""
    found = []
    for pat in FORBIDDEN_PATTERNS:
        if pat in text:
            found.append((source_name, pat))
    return found


def scan_file(file_path: Path) -> list:
    """Scans a single file (markdown, json, txt, docx)."""
    found = []
    suffix = file_path.suffix.lower()
    
    if suffix in [".md", ".json", ".txt", ".yaml", ".yml"]:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            found.extend(scan_text(content, str(file_path)))
        except UnicodeDecodeError:
            pass
    elif suffix == ".docx":
        try:
            doc = docx.Document(str(file_path))
            texts = [p.text for p in doc.paragraphs]
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        texts.append(c.text)
            combined = "\n".join(texts)
            found.extend(scan_text(combined, str(file_path)))
        except Exception as e:
            print(f"Warning: could not scan {file_path}: {e}", file=sys.stderr)
            
    return found


def scan_directory(dir_path: str) -> list:
    """Recursively scans a directory for leaked placeholders."""
    p = Path(dir_path).resolve()
    all_found = []
    for file in p.rglob("*"):
        if file.is_file() and file.suffix.lower() in [".md", ".json", ".txt", ".docx"]:
            all_found.extend(scan_file(file))
    return all_found


def main():
    parser = argparse.ArgumentParser(description="Scan rendered files for dangling placeholders.")
    parser.add_argument("target", help="File or directory to scan.")
    args = parser.parse_args()
    
    target_path = Path(args.target)
    if target_path.is_dir():
        results = scan_directory(args.target)
    else:
        results = scan_file(target_path)
        
    if results:
        print(f"FAILED: Found {len(results)} forbidden placeholder(s):", file=sys.stderr)
        for src, pat in results:
            print(f"  - [{pat}] in {src}", file=sys.stderr)
        sys.exit(1)
    else:
        print("PASSED: Zero dangling placeholders detected.")
        sys.exit(0)


if __name__ == "__main__":
    main()
